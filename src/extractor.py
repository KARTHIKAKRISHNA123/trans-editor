

import logging
from pathlib import Path
from docx import Document

# Set up logger for this module
# logging is better than print() in production:
# → can redirect to files, monitoring tools
# → has levels: DEBUG < INFO < WARNING < ERROR < CRITICAL
# → %(name)s shows "src.extractor" in the log line
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(name)s] %(levelname)s: %(message)s"
)
logger = logging.getLogger(__name__)



# ──────────────────────────────────────────────────────────
# Data Classes — snapshots of document structure
# ──────────────────────────────────────────────────────────

class RunData:
    
    def __init__(self, text: str, bold: bool, italic: bool, underline: bool):
        self.text      = text       # the text string of this run
        self.bold      = bold       # was it bold?    True/False
        self.italic    = italic     # was it italic?  True/False
        self.underline = underline  # was it underlined? True/False

    def __repr__(self):
        
        return f"RunData('{self.text[:20]}', bold={self.bold}, italic={self.italic})"


class ParagraphData:
    
    def __init__(self, para_ref, runs: list, full_text: str, alignment, style_name: str):
        self.para_ref   = para_ref    # live reference — needed for writing back
        self.runs       = runs        # List[RunData]
        self.full_text  = full_text   # complete English text of this paragraph
        self.alignment  = alignment   # python-docx WD_ALIGN_PARAGRAPH enum value
        self.style_name = style_name  # e.g., "Heading 1", "Normal"

    def __repr__(self):
        return f"ParagraphData('{self.full_text[:40]}', style='{self.style_name}')"


class TableData:
    
    def __init__(self, table_ref, cells: dict):
        self.table_ref = table_ref  # live python-docx Table reference
        self.cells     = cells      # dict: {(row, col): [ParagraphData, ...]}

    def __repr__(self):
        row_count = len(set(r for r, c in self.cells))
        col_count = len(set(c for r, c in self.cells))
        return f"TableData({row_count} rows × {col_count} cols)"


class DocumentData:
    def __init__(self, doc_ref, paragraphs: list, tables: list):
        self.doc_ref    = doc_ref    # live python-docx Document object
        self.paragraphs = paragraphs # List[ParagraphData]  — body paragraphs
        self.tables     = tables     # List[TableData]      — tables


# ──────────────────────────────────────────────────────────
# Private helper functions (underscore = internal use only)
# ──────────────────────────────────────────────────────────

def _extract_runs(para) -> tuple[list, str]:
    
    runs      = []
    full_text = ""

    for run in para.runs:
        # run.text is the actual string content of this run
        full_text += run.text

        runs.append(RunData(
            text      = run.text,
            bold      = bool(run.bold),       # None → False
            italic    = bool(run.italic),     # None → False
            underline = bool(run.underline),  # None → False
        ))

    return runs, full_text


def _extract_paragraph(para) -> ParagraphData:
    """
    Convert one python-docx Paragraph into a ParagraphData snapshot.
    """
    runs, full_text = _extract_runs(para)

    return ParagraphData(
        para_ref   = para,             # keep live reference for later writing
        runs       = runs,
        full_text  = full_text,
        alignment  = para.alignment,   # WD_ALIGN_PARAGRAPH or None
        style_name = para.style.name,  # "Normal", "Heading 1", etc.
    )


def _extract_table(table) -> TableData:
    """
    Convert one python-docx Table into a TableData snapshot.

    table.rows          → list of Row objects
    row.cells           → list of Cell objects
    cell.paragraphs     → list of Paragraph objects inside that cell
    """
    cells = {}

    for row_idx, row in enumerate(table.rows):
        # enumerate() gives us (0, row_obj), (1, row_obj), ... 
        for col_idx, cell in enumerate(row.cells):
            # each cell can have multiple paragraphs
            cell_paras = [_extract_paragraph(p) for p in cell.paragraphs]
            cells[(row_idx, col_idx)] = cell_paras

    return TableData(table_ref=table, cells=cells)


# ──────────────────────────────────────────────────────────
# Public API — the only function you call from outside
# ──────────────────────────────────────────────────────────

def extract_document(filepath: str) -> DocumentData:
    """
    Read a .docx file and return a DocumentData snapshot.

    Args:
        filepath : path to the .docx file

    Returns:
        DocumentData with all paragraphs and tables extracted

    Raises:
        FileNotFoundError : file doesn't exist
        ValueError        : file is not a .docx
        RuntimeError      : python-docx failed to parse the file

    Usage:
        data = extract_document("D:/docs/report.docx")
        print(data.paragraphs[0].full_text)   # first paragraph text
        print(data.tables[0])                  # first table summary
    """
    path = Path(filepath)

    # Guard: file must exist
    if not path.exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    # Guard: must be a .docx (not .doc, .pdf, .txt)
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Expected .docx, got: {path.suffix}")

    logger.info(f"Opening document: {filepath}")

    try:
        # Document() unzips the .docx and parses word/document.xml into objects
        doc = Document(str(path))
    except Exception as e:
        raise RuntimeError(f"Failed to open {filepath}: {e}")

    # Extract all body paragraphs
    # doc.paragraphs returns ALL paragraphs in the document body
    paragraphs = [_extract_paragraph(p) for p in doc.paragraphs]
    logger.info(f"Extracted {len(paragraphs)} paragraphs")

    # Extract all tables
    # doc.tables returns ALL tables in the document body
    tables = [_extract_table(t) for t in doc.tables]
    logger.info(f"Extracted {len(tables)} tables")

    return DocumentData(doc_ref=doc, paragraphs=paragraphs, tables=tables)